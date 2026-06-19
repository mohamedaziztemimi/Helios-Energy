"""
Word document generation tool. Produces a real .docx file using python-docx
— not mocked. Same input shape as the PDF tool (title, content, optional
table) so the agent can use either with the same mental model.
"""

import os
import re
import uuid

from langchain_core.tools import tool
from docx import Document


def _safe_filename(title: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9_-]+", "_", title.strip()).strip("_").lower()
    return f"{slug or 'report'}_{uuid.uuid4().hex[:8]}.docx"


def make_generate_word_tool(output_dir: str):
    """
    output_dir: the directory this run's generated files should be written
    to, e.g. /tmp/agent_outputs/<run_id>/.
    """

    @tool
    def generate_word(title: str, content: str, table_data: list[list[str]] | None = None) -> str:
        """
        Generate a real, downloadable Word (.docx) report. Returns the file path.

        title: the report's title, used as the document heading.
        content: the body text, in plain paragraphs (use \\n\\n between paragraphs).
        table_data: optional. A list of rows, where the first row is the
            header. Example: [["Plant", "Energy (kWh)"], ["Plant C1-001", "1234.5"]]
        """
        os.makedirs(output_dir, exist_ok=True)
        filename = _safe_filename(title)
        filepath = os.path.join(output_dir, filename)

        doc = Document()
        doc.add_heading(title, level=1)

        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        if table_data:
            rows, cols = len(table_data), len(table_data[0])
            table = doc.add_table(rows=1, cols=cols)
            table.style = "Light Grid Accent 1"

            header_cells = table.rows[0].cells
            for i, value in enumerate(table_data[0]):
                header_cells[i].text = str(value)

            for row_data in table_data[1:]:
                row_cells = table.add_row().cells
                for i, value in enumerate(row_data):
                    row_cells[i].text = str(value)

        doc.save(filepath)
        return filepath

    return generate_word